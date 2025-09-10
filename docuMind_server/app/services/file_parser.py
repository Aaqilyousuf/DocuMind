import io
import os
import PyPDF2
import docx

def extract_text(file_source, filename_hint: str = None) -> str:
    """
    Extract text from a file path or BytesIO object.

    Args:
        file_source: str (file path) or BytesIO object
        filename_hint: optional filename to detect extension (required if file_source is BytesIO)

    Returns:
        str: extracted text
    """
    # Determine extension
    if isinstance(file_source, str):
        # Local file path
        ext = os.path.splitext(file_source)[1].lower()
    elif isinstance(file_source, io.BytesIO):
        if not filename_hint:
            raise ValueError("filename_hint is required for file streams")
        ext = os.path.splitext(filename_hint)[1].lower()
    else:
        raise ValueError("file_source must be a file path or BytesIO object")

    # Route to correct extractor
    if ext == ".pdf":
        return extract_pdf(file_source)
    elif ext == ".docx":
        return extract_docx(file_source)
    elif ext in [".txt", ".csv"]:
        return extract_txt(file_source)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def extract_pdf(file_source) -> str:
    """Extract text from PDF (path or BytesIO)"""
    text = ""
    if isinstance(file_source, str):
        with open(file_source, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
    else:  # BytesIO
        reader = PyPDF2.PdfReader(file_source)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text


def extract_docx(file_source) -> str:
    """Extract text from DOCX (path or BytesIO)"""
    if isinstance(file_source, str):
        doc = docx.Document(file_source)
    else:
        doc = docx.Document(file_source)
    return "\n".join([para.text for para in doc.paragraphs])


def extract_txt(file_source) -> str:
    """Extract text from TXT or CSV (path or BytesIO)"""
    if isinstance(file_source, str):
        with open(file_source, "rb") as f:
            raw_bytes = f.read()
    else:
        file_source.seek(0)
        raw_bytes = file_source.read()
    
    try:
        return raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return raw_bytes.decode("utf-8", errors="ignore")
